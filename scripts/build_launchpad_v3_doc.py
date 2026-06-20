from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("E:/dapp") / "多零直发发射台"
OUT = ROOT / "多零直发持币分红V3_项目功能说明_修改版.docx"
TABLE_INDENT_DXA = 120


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.25):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_para(doc, text, size=11, bold=False, color="000000", before=0, after=6):
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, before=before, after=after)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    if level == 1:
        set_paragraph_spacing(paragraph, before=18, after=10)
        size, color = 16, "2E74B5"
    else:
        set_paragraph_spacing(paragraph, before=14, after=7)
        size, color = 13, "2E74B5"
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=True, color=color)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    set_paragraph_spacing(paragraph, after=4, line=1.25)
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    run = paragraph.add_run(text)
    set_run_font(run, size=11, color="000000")
    return paragraph


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, bold=False, fill=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    set_paragraph_spacing(paragraph, before=0, after=0, line=1.15)
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, bold=bold, color="000000")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    if fill:
        shade_cell(cell, fill)


def table_borders(table, color="D9DEE8"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)


def label_table(doc, rows, widths=(2500, 6860)):
    table = doc.add_table(rows=0, cols=2)
    set_table_geometry(table, widths)
    table_borders(table)
    for label, value in rows:
        cells = table.add_row().cells
        for cell, width in zip(cells, widths):
            set_cell_width(cell, width)
        set_cell_text(cells[0], label, bold=True, fill="E8EEF5")
        set_cell_text(cells[1], value)
    doc.add_paragraph()
    return table


def matrix_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    table_borders(table)
    for cell, header, width in zip(table.rows[0].cells, headers, widths):
        set_cell_width(cell, width)
        set_cell_text(cell, header, bold=True, fill="E8EEF5")
    for row in rows:
        cells = table.add_row().cells
        for cell, value, width in zip(cells, row, widths):
            set_cell_width(cell, width)
            set_cell_text(cell, value)
    doc.add_paragraph()
    return table


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("多零直发持币分红V3 项目功能说明")
    set_run_font(footer_run, size=9, color="666666")

    title = doc.add_paragraph()
    set_paragraph_spacing(title, before=0, after=3)
    title_run = title.add_run("多零直发持币分红V3")
    set_run_font(title_run, size=24, bold=True, color="0B2545")

    add_para(
        doc,
        "创建Token / 30个0（30个零）供应量 / 默认BNB底池 / 15%平台税 / 普通白名单与限购版名单",
        size=12,
        color="555555",
        after=14,
    )
    add_para(
        doc,
        "定位：这是一个偏傻瓜式的多零代币创建工具。项目方填写名称、符号和供应量后，即可创建带默认BNB底池、15%平台税收和白名单能力的代币。创建费用为0.005 BNB，后续由创建者钱包完成加池、开盘和运营管理。",
        after=10,
    )

    add_heading(doc, "一、项目定位", 1)
    label_table(
        doc,
        [
            ("项目类型", "多零直发代币创建工具，偏持币分红V3模式"),
            ("适用场景", "直接发行多0代币，创建后由项目钱包管理权限"),
            ("核心卖点", "支持0精度和30个0（30个零）供应量，默认BNB底池，默认写入15%平台税收，支持Pancake加池和手动开盘"),
            ("操作方式", "前端填写参数，一键上链创建，后续通过项目方钱包管理"),
            ("展示风格", "类似创建Token页面，左侧展示卖点，右侧填写参数、普通白名单和限购版名单选项"),
        ],
    )

    add_heading(doc, "二、发行参数", 1)
    label_table(
        doc,
        [
            ("代币名称", "项目方自定义"),
            ("代币符号", "项目方自定义"),
            ("初始供应量", "支持30个0（30个零）级别供应量，例如 1 后面 30 个 0 或 21 后面 30 个 0"),
            ("小数位", "建议0位小数，按整数发行，更适合多零代币"),
            ("交易对", "默认Pancake WBNB，也就是默认BNB底池"),
            ("底池币种", "默认BNB，不让用户复杂选择"),
            ("创建费用", "0.005 BNB，创建代币时支付"),
            ("开盘方式", "支持手动开启交易，建议先加池再开盘"),
            ("创建后权限", "谁创建代币，谁拥有该代币的管理权限"),
        ],
    )

    add_heading(doc, "三、自动税收模板", 1)
    add_para(
        doc,
        "平台创建的每个代币会自动带入默认税收模板，不需要项目方逐项复杂配置。默认所有交易税收按15%进入平台地址。税收在链上可查，前端可以做简化展示，业务上以合约参数和平台地址为准。",
        after=8,
    )
    matrix_table(
        doc,
        ["功能", "默认方向", "说明"],
        [
            ("平台税收", "默认15%", "所有税收默认进入平台地址，创建后自动生效。"),
            ("持币分红", "按版本配置", "保留持币分红V3能力，可按持币比例分配USDT或指定分红币。"),
            ("销毁", "按版本配置", "可保留销毁逻辑，减少流通量。"),
            ("回流加池", "按版本配置", "可保留回流LP逻辑，增强池子深度。"),
            ("税收展示", "前端简化", "页面可减少复杂字段，但链上规则仍以合约参数为准。"),
        ],
        [2100, 1600, 5660],
    )

    add_heading(doc, "四、白名单机制", 1)
    add_para(
        doc,
        "白名单拆成两个选项：普通白名单和限购版名单，避免项目方地址、做市地址和用户限购名单混在一起。",
        after=8,
    )
    matrix_table(
        doc,
        ["白名单类型", "使用对象", "作用"],
        [
            ("普通白名单", "项目方、加池钱包、做市钱包、测试钱包", "用于开盘前转账、加池、测试和必要管理，可按合约配置免交易限制。"),
            ("限购版名单", "限购阶段用户", "只给买入资格和额度，不给管理权限；仍受单笔上限、钱包上限或白名单份数限制。"),
        ],
        [2100, 2800, 4460],
    )
    add_bullet(doc, "普通白名单可以一次添加多个地址，适合项目方批量导入。")
    add_bullet(doc, "限购版名单可以设置每个地址可买份数或最大买入量，用于限购阶段预热和控盘。")
    add_bullet(doc, "普通白名单和限购版名单互相独立，避免用户拿到项目方级别权限。")

    add_heading(doc, "五、创建与开盘流程", 1)
    add_bullet(doc, "填写代币名称、符号、供应量，默认支持30个0级别供应量。")
    add_bullet(doc, "创建时支付0.005 BNB创建费用。")
    add_bullet(doc, "系统自动套用15%平台税收模板，税收进入平台地址。")
    add_bullet(doc, "交易对默认Pancake WBNB，也就是默认BNB底池。")
    add_bullet(doc, "按需要添加普通白名单和限购版名单。")
    add_bullet(doc, "创建代币后先完成授权和加池。")
    add_bullet(doc, "确认LP和参数无误后，由创建者钱包手动开盘。")

    add_heading(doc, "六、页面功能建议", 1)
    matrix_table(
        doc,
        ["页面模块", "建议字段", "说明"],
        [
            ("创建Token", "名称、符号、供应量、默认BNB底池", "保持简单，默认值直接给30个0供应量和15%平台税收。"),
            ("税收设置", "平台地址、买入税、卖出税", "默认15%进入平台地址，高级参数可以折叠，避免页面太乱。"),
            ("普通白名单", "多地址输入、批量添加、批量删除", "给项目方和加池相关钱包使用。"),
            ("限购版名单", "地址、额度、钱包上限、限购份数", "给用户限购阶段使用。"),
            ("开盘管理", "加池状态、交易状态、手动开盘按钮", "开盘前提示先加池，减少误操作。"),
        ],
        [1800, 3300, 4260],
    )

    add_heading(doc, "七、最终确认点", 1)
    add_bullet(doc, "所有通过平台创建的币，都自动带默认15%平台税收模板。")
    add_bullet(doc, "默认BNB底池，交易对使用Pancake WBNB。")
    add_bullet(doc, "创建费用为0.005 BNB。")
    add_bullet(doc, "支持30个0级别供应量，不限制只能15个0。")
    add_bullet(doc, "普通白名单支持多个地址批量管理。")
    add_bullet(doc, "限购版名单是单独选项，只控制买入资格和额度。")
    add_bullet(doc, "创建者就是该代币管理员，不再把权限留给平台部署钱包。")
    add_bullet(doc, "手动开盘保留，开盘前先完成加池和参数检查。")

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build())
