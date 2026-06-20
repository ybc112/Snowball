from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("E:/dapp") / "多零直发发射台"
OUT = ROOT / "杀0协议_项目功能说明_税收0.6回流黑洞版.docx"

PAGE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.25):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_para(doc, text="", size=11, bold=False, color="000000", after=6, before=0):
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, before=before, after=after)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    if level == 1:
        set_paragraph_spacing(paragraph, before=18, after=10)
        size, color = 16, "2E74B5"
    else:
        set_paragraph_spacing(paragraph, before=14, after=7)
        size, color = 13, "2E74B5"
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=True, color=color)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    set_paragraph_spacing(paragraph, after=4, line=1.25)
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    if paragraph.runs:
        paragraph.runs[0].text = text
        set_run_font(paragraph.runs[0], size=11, color="000000")
    else:
        run = paragraph.add_run(text)
        set_run_font(run, size=11, color="000000")
    return paragraph


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, bold=False, fill=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    set_paragraph_spacing(paragraph, before=0, after=0, line=1.15)
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, bold=bold, color="000000")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    if fill:
        shade_cell(cell, fill)


def table_borders(table, color="D9DEE8"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)


def label_table(doc, rows, widths=(2700, 6660)):
    table = doc.add_table(rows=0, cols=2)
    set_table_geometry(table, widths)
    table_borders(table)
    for label, value in rows:
        cells = table.add_row().cells
        for cell, width in zip(cells, widths):
            set_cell_width(cell, width)
        set_cell_text(cells[0], label, bold=True, fill="E8EEF5")
        set_cell_text(cells[1], value)
    doc.add_paragraph()
    return table


def matrix_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    table_borders(table)
    for cell, header, width in zip(table.rows[0].cells, headers, widths):
        set_cell_width(cell, width)
        set_cell_text(cell, header, bold=True, fill="E8EEF5")
    for row in rows:
        cells = table.add_row().cells
        for cell, value, width in zip(cells, row, widths):
            set_cell_width(cell, width)
            set_cell_text(cell, value)
    doc.add_paragraph()
    return table


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("杀0协议 项目功能说明")
    set_run_font(footer_run, size=9, color="666666")

    title = doc.add_paragraph()
    set_paragraph_spacing(title, before=0, after=3)
    title_run = title.add_run("杀0协议")
    set_run_font(title_run, size=24, bold=True, color="0B2545")

    add_para(
        doc,
        "多零直发代币 / USDT持币分红 / 手动开盘 / 支持30个零加LP",
        size=12,
        color="555555",
        after=14,
    )

    add_para(
        doc,
        "定位：单独发行一个多零代币合约，不走复杂发射台认购流程。项目方先部署合约，再配置白名单、税收、分红和开盘参数，完成LP后手动开盘。",
        after=10,
    )

    add_heading(doc, "一、核心参数", 1)
    label_table(
        doc,
        [
            ("项目名称", "杀0协议"),
            ("合约类型", "单独发行的多零代币合约"),
            ("总供应量", "21 后面 30 个 0，即 21 x 10^30"),
            ("小数位", "0 位小数，链上按整数数量计算，适合多零直发"),
            ("靓号要求", "Token 合约地址尾号 6 个 0"),
            ("分红币种", "USDT 持币分红"),
            ("开盘方式", "手动开盘：先加LP，再由项目方开启交易"),
            ("白名单", "保留白名单权限，方便开盘前加池、测试和项目方操作"),
            ("LP重点", "支持 30 个 0 级别供应量正常授权并加入 Pancake LP"),
        ],
    )

    add_heading(doc, "二、税收机制", 1)
    add_para(
        doc,
        "买卖税收由项目方按需求设置。以总税 3% 为例，项目/营销钱包拿的是所有税收的 20%，不是交易额的 20%；换算后为交易额的 0.6%，按直达钱包逻辑进入指定地址。其余税收按合约配置执行，其中销毁 0.5%，回流 0.5%；回流用于加池，产生的 LP 进入黑洞地址，避免后续被取回。链上税收无法真正隐藏，源码和交易记录都能查到；前端展示可以简化。",
        after=8,
    )
    matrix_table(
        doc,
        ["项目", "比例", "说明"],
        [
            ("示例总税", "3%", "买入和卖出可按项目方设置计算；这里用3%作为示例。"),
            ("项目/营销钱包", "税收的20% / 交易额0.6%", "不是额外20%税，而是示例3%税收里的20%，交易触发后直达指定钱包。"),
            ("剩余税收", "税收的80% / 交易额2.4%", "按合约配置继续进入持币分红、销毁、回流等逻辑。"),
            ("销毁", "0.5%", "按交易税逻辑进入销毁流程，减少流通压力。"),
            ("回流", "0.5%", "用于自动回流加池，回流产生的 LP 进入黑洞地址。"),
            ("持币分红", "剩余配置", "剩余部分进入分红逻辑，按持币规则发放USDT。"),
        ],
        [2200, 1500, 5660],
    )

    add_heading(doc, "三、到账和触发方式", 1)
    add_bullet(doc, "项目/营销钱包：拿项目方设置税收的20%；以总税3%为例，等于交易额0.6%，交易触发后直达指定地址。")
    add_bullet(doc, "剩余税收：默认占总税收80%，即交易额2.4%，继续按合约配置进入分红、销毁、回流。")
    add_bullet(doc, "销毁：默认0.5%，按交易税逻辑进入销毁流程，减少流通压力。")
    add_bullet(doc, "回流：默认0.5%，用于自动回流加池；回流产生的LP进入黑洞地址，不进项目方钱包。")
    add_bullet(doc, "持币分红：剩余部分进入分红池，达到处理条件后按持币比例分发USDT。")
    add_bullet(doc, "如果采用原参考合约逻辑，税会先进入合约，卖出时再触发自动处理；如果要“直接到账”，需要使用改过的新版合约。")

    add_heading(doc, "四、参考合约方向", 1)
    label_table(
        doc,
        [
            ("参考合约A", "0xdbb64ec2a453a612d0cc5001b6f15b0200000000"),
            ("参考重点", "0位小数、多零总量、买卖税、持币分红、白名单、手动开盘、自动处理。"),
            ("参考合约B", "0x2765e179b2a0bc0828ad83f3570ebd62b9e00000"),
            ("参考重点", "靓号尾号方案；客户反馈该方向偏向15个0加LP，主方案仍按A的30个0能力处理。"),
        ],
    )

    add_heading(doc, "五、玩法流程", 1)
    add_bullet(doc, "部署杀0协议代币，供应量设置为 21 后面 30 个 0。")
    add_bullet(doc, "把项目方、加池钱包、测试钱包加入白名单。")
    add_bullet(doc, "配置USDT分红、项目税钱包、销毁和回流参数。")
    add_bullet(doc, "授权Token并添加Pancake LP，确认30个零数量可以正常进池。")
    add_bullet(doc, "LP完成后手动开盘，正式开放买卖。")

    add_heading(doc, "六、需要保留的权限", 1)
    label_table(
        doc,
        [
            ("白名单权限", "保留，用于开盘前配置、测试和特殊地址放行。"),
            ("手动开盘权限", "保留，LP完成后由项目方开启交易。"),
            ("税收钱包", "部署前设置项目/营销税收款地址。"),
            ("分红设置", "保留USDT分红参数，方便后续调整最低持币和处理阈值。"),
            ("空投裂变", "默认3，用于基础空投裂变展示和转账扩散；后续可根据Gas和项目需求调整。"),
        ],
    )

    add_heading(doc, "七、最终确认点", 1)
    add_bullet(doc, "合约必须支持 21 后面 30 个 0 的总量。")
    add_bullet(doc, "30个零数量必须可以正常Approve和Add LP。")
    add_bullet(doc, "项目/营销钱包拿项目方设置税收的20%；以总税3%为例，实际为交易额0.6%，并做成直达指定钱包。")
    add_bullet(doc, "销毁默认0.5%，回流默认0.5%，回流加池产生的LP进入黑洞地址。")
    add_bullet(doc, "持币分红使用USDT，剩余税收按分红和其他配置执行。")
    add_bullet(doc, "总税、分红、销毁、回流、黑洞地址等参数需部署前确认。")
    add_bullet(doc, "保留白名单和手动开盘，空投裂变默认3个地址。")

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build())
